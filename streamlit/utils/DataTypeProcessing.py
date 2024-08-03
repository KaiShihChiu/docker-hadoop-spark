import io
from pdfminer.high_level import extract_text_to_fp
from pdf2image import convert_from_bytes
from PIL import Image, ImageEnhance
import pytesseract
import os
import subprocess
from docx import Document


class PdfProcessor:
    def __init__(self, crop_ratio=(0.1, 0.05, 0.9, 0.9)):
        self.crop_ratio = crop_ratio
    
    def extract_text(self, content):
        output_string = io.StringIO()
        with io.BytesIO(content) as open_pdf_file:
            extract_text_to_fp(open_pdf_file, output_string)
        return output_string.getvalue()
    
    def preprocess_image(self, image):
        # 灰度化
        image = image.convert('L')
        width, height = image.size
        # 按比例裁剪图片
        left = width * self.crop_ratio[0]
        upper = height * self.crop_ratio[1]
        right = width * self.crop_ratio[2]
        lower = height * self.crop_ratio[3]
        image = image.crop((left, upper, right, lower))
        # 增强对比度
        enhancer = ImageEnhance.Contrast(image)
        image = enhancer.enhance(2)
        # 二值化
        image = image.point(lambda x: 0 if x < 220 else 255, '1')
        return image
    
    def ocr_images(self, content):
        images = convert_from_bytes(content)
        text_content = ""
        for image in images:
            preprocessed_image = self.preprocess_image(image)
            text = pytesseract.image_to_string(preprocessed_image, lang='chi_tra')
            text_content += text + "\n"
        return text_content

    def process_pdf(self, content):
        text_content = self.extract_text(content)
        if len(text_content.strip()) > 0:
            # PDF 文件主要是文字
            return text_content, False
        else:
            # PDF 文件主要是图片
            text_content = self.ocr_images(content)
            print(text_content)
            return text_content, True
        
from docx import Document


class DocProcessor:
    def __init__(self):
        pass

    def convert_doc_to_docx(self, doc_content, filename):
        # Save the uploaded doc file to a temporary location
        temp_doc_path = os.path.join(os.getcwd(), filename)
        with open(temp_doc_path, 'wb') as f:
            f.write(doc_content)

        # Convert .doc to .docx using libreoffice
        subprocess.run(['libreoffice', '--headless', '--convert-to', 'docx', temp_doc_path], stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        
        docx_path = temp_doc_path.replace('.doc', '.docx')
        return docx_path, temp_doc_path

    def read_all_text(self, docx_path):
        doc = Document(docx_path)
        all_text = []

        # Extract text from paragraphs
        for para in doc.paragraphs:
            all_text.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text.append(cell.text.strip())

        return "\n".join(all_text)

    def remove_duplicate_sentences(self, text):
        sentences = text.split("\n")
        unique_sentences = list(dict.fromkeys(sentences))
        return "\n".join(unique_sentences)
    
    
    def process_doc(self, content, filename):
        docx_path, temp_doc_path = self.convert_doc_to_docx(content, filename)

        if docx_path:
            all_text = self.read_all_text(docx_path)
            unique_text = self.remove_duplicate_sentences(all_text)
            
            # Clean up temporary files
            os.remove(docx_path)
            os.remove(temp_doc_path)
        return unique_text

class DocxProcessor:
    def __init__(self):
        pass

    def extract_text_from_docx(self, docx_content):
        # 使用 BytesIO 将二进制内容转换为一个文件对象
        docx_file = io.BytesIO(docx_content)
        # 使用 python-docx 读取文件内容
        document = Document(docx_file)
        text = "\n".join([para.text for para in document.paragraphs])
        return text

    def process_docx(self, docx_content):
        text_content = self.extract_text_from_docx(docx_content)
        return text_content


import pytesseract
from PIL import Image, ImageEnhance, ImageFilter
import cv2
import numpy as np

class ImgProcessor:
    def __init__(self, image_bytes):
        self.image = Image.open(io.BytesIO(image_bytes))
    
    def convert_to_grayscale(self):
        self.image = self.image.convert('L')
        return self.image

    def enhance_contrast(self, factor=2):
        enhancer = ImageEnhance.Contrast(self.image)
        self.image = enhancer.enhance(factor)
        return self.image

    def upsample_image(self, scale_percent=200):
        image_np = np.array(self.image)
        width = int(image_np.shape[1] * scale_percent / 100)
        height = int(image_np.shape[0] * scale_percent / 100)
        dim = (width, height)
        self.image = cv2.resize(image_np, dim, interpolation=cv2.INTER_CUBIC)
        return self.image

    def sharpen_image(self):
        kernel = np.array([[0, -1, 0], 
                           [-1, 5, -1], 
                           [0, -1, 0]])
        self.image = cv2.filter2D(self.image, -1, kernel)
        return self.image

    def adaptive_threshold(self):
        if len(self.image.shape) == 2 or self.image.shape[2] == 1:
            self.image = cv2.cvtColor(self.image, cv2.COLOR_GRAY2BGR)
        gray = cv2.cvtColor(self.image, cv2.COLOR_BGR2GRAY)
        self.binary = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_MEAN_C, cv2.THRESH_BINARY_INV, 15, 9)
        return self.binary

    def remove_horizontal_lines(self):
        horizontal_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (15, 1))
        detect_horizontal = cv2.morphologyEx(self.binary, cv2.MORPH_OPEN, horizontal_kernel, iterations=2)
        contours, _ = cv2.findContours(detect_horizontal, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        for contour in contours:
            [x, y, w, h] = cv2.boundingRect(contour)
            if w > 120:
                cv2.drawContours(self.image, [contour], -1, (255, 255, 255), thickness=cv2.FILLED)
        return self.image

    def extract_text(self, lang='chi_tra'):
        self.image = self.convert_to_grayscale()
        self.image = self.enhance_contrast()
        self.image = self.upsample_image()
        self.image = self.sharpen_image()
        self.binary = self.adaptive_threshold()
        self.image = self.remove_horizontal_lines()
        
        text = pytesseract.image_to_string(self.image, lang=lang)
        return text


# 示例使用
if __name__ == '__main__':
    content = b'%PDF-1.4...'  # PDF 文件的二进制内容
    pdf_processor = PdfProcessor()

    text_content, show_col2 = pdf_processor.process_pdf(content)
    if show_col2:
        st.session_state['show_col2'] = True
    else:
        st.session_state['show_col2'] = False
