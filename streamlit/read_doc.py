import streamlit as st
import subprocess
import os
from docx import Document

def convert_doc_to_docx(doc_content, filename):
    # Save the uploaded doc file to a temporary location
    temp_doc_path = os.path.join(os.getcwd(), filename)
    with open(temp_doc_path, 'wb') as f:
        f.write(doc_content)

    # Convert .doc to .docx using libreoffice
    subprocess.run(['libreoffice', '--headless', '--convert-to', 'docx', temp_doc_path], stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    
    docx_path = temp_doc_path.replace('.doc', '.docx')
    return docx_path, temp_doc_path

def read_all_text(docx_path):
    doc = Document(docx_path)
    all_text = []

    # Extract text from paragraphs
    for para in doc.paragraphs:
        all_text.append(para.text)
    
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text.append(cell.text.strip())

    return "\n".join(all_text)

def remove_duplicate_sentences(text):
    sentences = text.split("\n")
    unique_sentences = list(dict.fromkeys(sentences))
    return "\n".join(unique_sentences)

st.title("Read DOC File Content and Remove Duplicate Sentences")

uploaded_file = st.file_uploader("Upload a DOC file", type="doc")

if uploaded_file is not None:
    st.write("File uploaded successfully.")
    docx_path, temp_doc_path = convert_doc_to_docx(uploaded_file.read(), uploaded_file.name)

    if docx_path:
        all_text = read_all_text(docx_path)
        unique_text = remove_duplicate_sentences(all_text)
        
        # Display all text content without duplicates
        st.header("Document Text Content")
        st.text_area("Content", unique_text, height=400)
        
        # Clean up temporary files
        os.remove(docx_path)
        os.remove(temp_doc_path)
