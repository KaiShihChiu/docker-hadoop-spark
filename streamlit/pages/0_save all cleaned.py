from pathlib import Path
import sys
import os
from hdfs import InsecureClient
from pdf2image import convert_from_path
import pytesseract
import io
from PIL import Image, ImageEnhance, ImageFilter
from docx import Document
current_dir = os.path.dirname(os.path.abspath(__file__))
import sys
sys.path.append(current_dir)
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'utils')))
from pdfminer.high_level import extract_text, extract_text_to_fp
from DataCleaner import DataCleaner
from llm import improve_text
from DataTypeProcessing import PdfProcessor, DocxProcessor, DocProcessor, ImgProcessor
import pandas as pd
import streamlit as st
from hivedb import HiveDB

# 使用 NameNode 的主机名和端口号，确保这些在容器网络中可达
hdfs_url = 'http://namenode:9870'
hdfs_user = 'hadoop'  # 替换为您的 HDFS 用户名
fp = '/hadoop-data/'

# 初始化 HDFS 客户端
client = InsecureClient(hdfs_url, user=hdfs_user)
data_cleaner = DataCleaner()

# 列出 HDFS 中的文件
file_list = client.list('/hadoop-data')
# print(f"Files in HDFS: {file_list}")


# Initialize session state for tracking selected database, table, and table list
if 'selected_db' not in st.session_state:
    st.session_state.selected_db = None
if 'selected_table' not in st.session_state:
    st.session_state.selected_table = None
if 'tables' not in st.session_state:
    st.session_state.tables = ['Select a database first']

hive_db = HiveDB()
# Fetch and display databases
databases = hive_db.show_databases()
databases = ['Please select a database'] + databases
select_db = st.sidebar.selectbox("Select the database", databases)

# Check if the selected database has changed
if select_db != 'Please select a database' and select_db != st.session_state.selected_db:
    st.session_state.selected_db = select_db

    # Use the selected database
    hive_db.use_database(select_db)
    tables = hive_db.show_tables()
    st.session_state.tables = ['Please select a table'] + tables  # Reset table selection when database changes

# Fetch and display tables in the selected database
select_tb = st.sidebar.selectbox("Select the table to DELETE", st.session_state.tables)

if select_db != 'Please select a database' and select_tb != 'Please select a table':
    # add a delete button
    if st.sidebar.button("Delete"):
        hive_db.delete_table(select_tb)
        st.success(f"Successfully deleted the table: {select_tb}")


def process_selected_file(selected_file):
    # determine the file type
    file_type = selected_file.split('.')[-1]
    # print(f"File type: {file_type}")

    # 讀取文件內容
    hdfs_path = '/hadoop-data/' + selected_file

    with client.read(hdfs_path) as reader:
        content = reader.read()  # 以二進制模式讀取

    if file_type == 'pdf':
        pdf_processor = PdfProcessor()
        text_content, _ = pdf_processor.process_pdf(content)
        cleaned = data_cleaner.data_cleaning(text_content)

    elif file_type == 'docx':
        docx_processor  = DocxProcessor()
        docx_content = docx_processor.extract_text_from_docx(content)
        cleaned = data_cleaner.data_cleaning(docx_content)
            
    elif file_type in ['png', 'jpg', 'jpeg']:
        img_processor = ImgProcessor(content)
        text = img_processor.extract_text()
        cleaned = data_cleaner.data_cleaning(text)

    elif file_type == 'doc':
        doc_processor  = DocProcessor()
        doc_content = doc_processor.process_doc(content, selected_file)
        cleaned = data_cleaner.data_cleaning(doc_content)

    return cleaned

# monitor the data processing percentage using streamlit
st.write("Data processing progress")


# 设置进度条
progress_bar = st.progress(0)
status_text = st.empty()
# 模拟数据处理
# add a button to start the progress

# Initialize session state for stopping process
if 'stop' not in st.session_state:
    st.session_state.stop = False

def stop_progress():
    st.session_state.stop = True
    
# remove the file that is not ended with .pdf, .docx, .doc, .png, .jpg, .jpeg

df = pd.DataFrame(columns=['filename', 'cleaned'])
if st.button('Start Progress'):
    for i, selec_file in enumerate(file_list):
        if not selec_file.endswith(('.pdf', '.docx', '.doc', '.png', '.jpg', '.jpeg')):
            continue
        else:
            pass

        new_data = pd.DataFrame({
            'filename': [selec_file],
            'cleaned': [process_selected_file(selec_file)]
        })


        
        df = pd.concat([df, new_data], ignore_index=True)

        # 更新进度条
        progress = int((i + 1) / (len(file_list)) * 100)
        progress_bar.progress(progress)
        status_text.text(f"Progress: {progress}%")
        
    # st.write(df)
    # successfully cleaned!!
    st.success("Successfully cleaned the data!")
    
    # save the cleaned data to the database and show the spinned animation until the data is saved
    st.spinner("Saving the data to the database")
    # print(df)
    # check hive to see if the file is already there
    
    from hivedb import HiveDB
    hive_db = HiveDB()

    # # 显示所有数据库
    databases = hive_db.show_databases()
    st.write("old Databases:", databases)
    
    db = "gov_docs"
    # use the test database
    hive_db.use_database(db)
    databases = hive_db.show_databases()
    st.write("new Databases:", databases)
    
    # # write a df with multiple rows to hive
    # show table name in db
    old_tables = hive_db.show_tables()
    
    table_name = 'content_table'

    columns = hive_db.extract_colname_dtype(df)
    
    hive_db.insert_data_from_df(table_name, df)
    
    # # read the table
    st.write(hive_db.read_table(table_name))
    
    # show the success message
    st.success("Successfully saved the data to the database")


    # create a new table for the metadata
    metadata_table = 'metadata_table'
    metadata_columns = {
        'filename': 'STRING',
        'keywords': 'STRING'
    }
    
    





# selec_table = st.sidebar.selectbox("Select a table", tables)

# save the selected file in the table for future use


# # check if the file is already in the database
# file_exists = hive.check_file_exists(selected_file)
# if not file_exists:
#     # insert the file into the database
#     hive.insert_file(selected_file, cleaned)
        

    
# elif main_page == datatype[2]:
# from MetaDataExtractor import MetaDataExtractor
# metadata_extractor = MetaDataExtractor()

# print(st.session_state.cleaned)
# print(st.session_state.selected_file)
# metadata = metadata_extractor.extract_metadata(st.session_state.selected_file, st.session_state.cleaned)
# st.write("Metadata extraction")
